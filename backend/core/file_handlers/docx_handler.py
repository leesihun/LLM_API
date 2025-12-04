"""
DOCX File Handler
==================
Unified handler for Word .docx files.

Version: 2.0.0
Created: December 3, 2025
"""

from pathlib import Path
from typing import Dict, Any, List
from backend.core.file_handlers.base import FileHandler
from backend.utils.logging_utils import get_logger

logger = get_logger(__name__)


class DOCXHandler(FileHandler):
    """Handler for DOCX files"""

    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
        self.file_type = 'docx'

    def extract_metadata(
        self,
        file_path: Path,
        quick_mode: bool = False
    ) -> Dict[str, Any]:
        """Extract DOCX metadata for code generation"""
        metadata = {
            'file_type': 'docx',
            'file_size': self._get_file_size(file_path),
            'file_size_human': self._format_file_size(self._get_file_size(file_path)),
            'error': None
        }

        try:
            from docx import Document
            doc = Document(file_path)

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            total_text = '\n'.join(paragraphs)

            metadata['paragraphs'] = len(paragraphs)
            metadata['words'] = len(total_text.split())
            metadata['characters'] = len(total_text)
            metadata['tables'] = len(doc.tables)

            if not quick_mode:
                metadata['preview'] = total_text[:500]

        except ImportError:
            metadata['error'] = "python-docx not installed"
        except Exception as e:
            metadata['error'] = str(e)

        return metadata

    def analyze(self, file_path: Path) -> Dict[str, Any]:
        """Comprehensive DOCX analysis"""
        result = {
            'format': 'DOCX',
            'file_path': str(file_path.absolute()),
            'filename': file_path.name,
            'size_bytes': self._get_file_size(file_path),
            'size_human': self._format_file_size(self._get_file_size(file_path))
        }

        try:
            from docx import Document
            doc = Document(file_path)

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            total_text = '\n'.join(paragraphs)

            result.update({
                'total_paragraphs': len(paragraphs),
                'total_words': len(total_text.split()),
                'total_characters': len(total_text),
                'total_tables': len(doc.tables),
                'text_preview': total_text[:500]
            })

        except ImportError:
            result['error'] = "python-docx not installed"
        except Exception as e:
            result['error'] = str(e)

        return result
