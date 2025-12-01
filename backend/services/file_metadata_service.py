"""
File Metadata Service
=====================
Unified service for extracting metadata from various file formats.

Supported formats: CSV, Excel, JSON, Text, PDF, DOCX, Images

Version: 1.0.0
Created: 2025-01-13
"""

import json
from pathlib import Path
from typing import Dict, Any, Optional, List
import logging

from backend.utils.logging_utils import get_logger

logger = get_logger(__name__)


class FileMetadataService:
    """
    Service for extracting metadata from files.

    Provides unified interface for file metadata extraction across different formats.
    Used by python_coder_tool, file_analyzer_tool, and other components.

    Features:
    - Support for CSV, Excel, JSON, Text, PDF, DOCX, Images
    - Standardized metadata structure
    - Error handling and logging
    - Configurable depth for nested structures
    """

    def __init__(self):
        """Initialize the file metadata service."""
        self.supported_formats = {
            ".csv": self._extract_csv_metadata,
            ".tsv": self._extract_csv_metadata,
            ".xlsx": self._extract_excel_metadata,
            ".xls": self._extract_excel_metadata,
            ".xlsm": self._extract_excel_metadata,
            ".json": self._extract_json_metadata,
            ".txt": self._extract_text_metadata,
            ".md": self._extract_text_metadata,
            ".log": self._extract_text_metadata,
            ".rtf": self._extract_text_metadata,
            ".pdf": self._extract_pdf_metadata,
            ".docx": self._extract_docx_metadata,
            ".png": self._extract_image_metadata,
            ".jpg": self._extract_image_metadata,
            ".jpeg": self._extract_image_metadata,
            ".gif": self._extract_image_metadata,
            ".bmp": self._extract_image_metadata,
        }

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a file.

        Args:
            file_path: Path to the file

        Returns:
            Dictionary containing file metadata

        Example:
            >>> service = FileMetadataService()
            >>> metadata = service.extract_metadata("/path/to/file.csv")
            >>> print(metadata["type"], metadata["columns"])
        """
        path = Path(file_path)

        # Basic file info
        metadata = {
            "filename": path.name,
            "extension": path.suffix.lower(),
            "size_mb": round(path.stat().st_size / (1024 * 1024), 2),
            "exists": path.exists(),
        }

        if not path.exists():
            metadata["error"] = "File not found"
            return metadata

        # Extract format-specific metadata
        extension = path.suffix.lower()
        if extension in self.supported_formats:
            try:
                extractor = self.supported_formats[extension]
                format_metadata = extractor(path)
                metadata.update(format_metadata)
            except Exception as e:
                logger.warning(f"[FileMetadataService] Failed to extract metadata for {path.name}: {e}")
                metadata["error"] = str(e)
        else:
            metadata["type"] = "unsupported"
            metadata["note"] = f"Format '{extension}' not supported"

        return metadata

    def _extract_csv_metadata(self, path: Path) -> Dict[str, Any]:
        """
        Extract metadata from CSV/TSV files.

        Args:
            path: Path to CSV file

        Returns:
            Metadata dictionary with columns, dtypes, sample rows
        """
        try:
            import pandas as pd

            df = pd.read_csv(path, nrows=5)
            return {
                "type": "csv",
                "columns": df.columns.tolist(),
                "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                "sample_rows": len(df),
                "row_count": len(df)
            }
        except Exception as e:
            logger.warning(f"[FileMetadataService] CSV extraction error: {e}")
            return {"type": "csv", "error": str(e)}

    def _extract_excel_metadata(self, path: Path) -> Dict[str, Any]:
        """
        Extract metadata from Excel files.

        Args:
            path: Path to Excel file

        Returns:
            Metadata dictionary with sheets, columns, dtypes
        """
        try:
            import pandas as pd

            # Get sheet names
            excel_file = pd.ExcelFile(path)
            sheet_names = excel_file.sheet_names

            # Analyze first sheet
            df = pd.read_excel(path, nrows=5)

            return {
                "type": "excel",
                "sheets": sheet_names,
                "sheet_count": len(sheet_names),
                "columns": df.columns.tolist(),
                "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                "sample_rows": len(df)
            }
        except Exception as e:
            logger.warning(f"[FileMetadataService] Excel extraction error: {e}")
            return {"type": "excel", "error": str(e)}

    def _extract_json_metadata(self, path: Path) -> Dict[str, Any]:
        """
        Extract metadata from JSON files with deep structure analysis.

        Args:
            path: Path to JSON file

        Returns:
            Metadata dictionary with structure, keys, depth analysis
        """
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            metadata = {
                "type": "json",
                "structure": type(data).__name__,
            }

            # Structure-specific analysis
            if isinstance(data, list):
                metadata["items_count"] = len(data)
                if len(data) > 0:
                    metadata["first_item_type"] = type(data[0]).__name__
                    if isinstance(data[0], dict):
                        metadata["keys"] = list(data[0].keys())
                    # Safe preview (limit size)
                    metadata["preview"] = self._safe_json_preview(data[:3])
            elif isinstance(data, dict):
                metadata["keys"] = list(data.keys())
                metadata["preview"] = self._safe_json_preview(data)
            else:
                metadata["value"] = str(data)[:200]

            # Calculate depth
            metadata["max_depth"] = self._calculate_json_depth(data)

            return metadata

        except json.JSONDecodeError as e:
            return {"type": "json", "error": f"Invalid JSON: {e}"}
        except Exception as e:
            logger.warning(f"[FileMetadataService] JSON extraction error: {e}")
            return {"type": "json", "error": str(e)}

    def _extract_text_metadata(self, path: Path) -> Dict[str, Any]:
        """
        Extract metadata from text files.

        Args:
            path: Path to text file

        Returns:
            Metadata dictionary with line count and preview
        """
        try:
            with open(path, 'r', encoding='utf-8') as f:
                lines = f.readlines()

            return {
                "type": "text",
                "line_count": len(lines),
                "preview": ''.join(lines[:5])[:200]
            }
        except Exception as e:
            logger.warning(f"[FileMetadataService] Text extraction error: {e}")
            return {"type": "text", "error": str(e)}

    def _extract_pdf_metadata(self, path: Path) -> Dict[str, Any]:
        """
        Extract metadata from PDF files.

        Args:
            path: Path to PDF file

        Returns:
            Metadata dictionary with page count and text preview
        """
        try:
            import PyPDF2

            with open(path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                page_count = len(pdf.pages)

                # Extract first page text as preview
                preview = ""
                if page_count > 0:
                    first_page = pdf.pages[0]
                    preview = first_page.extract_text()[:200]

            return {
                "type": "pdf",
                "page_count": page_count,
                "preview": preview
            }
        except Exception as e:
            logger.warning(f"[FileMetadataService] PDF extraction error: {e}")
            return {"type": "pdf", "error": str(e)}

    def _extract_docx_metadata(self, path: Path) -> Dict[str, Any]:
        """
        Extract metadata from DOCX files.

        Args:
            path: Path to DOCX file

        Returns:
            Metadata dictionary with paragraph count, word count, preview
        """
        try:
            from docx import Document

            doc = Document(path)

            # Count paragraphs and words
            paragraph_count = len(doc.paragraphs)
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)

            # Get preview text
            preview = "\n".join(p.text for p in doc.paragraphs[:3])[:200]

            # Count tables
            table_count = len(doc.tables)

            return {
                "type": "docx",
                "paragraph_count": paragraph_count,
                "word_count": word_count,
                "table_count": table_count,
                "preview": preview
            }
        except Exception as e:
            logger.warning(f"[FileMetadataService] DOCX extraction error: {e}")
            return {"type": "docx", "error": str(e)}

    def _extract_image_metadata(self, path: Path) -> Dict[str, Any]:
        """
        Extract metadata from image files.

        Args:
            path: Path to image file

        Returns:
            Metadata dictionary with dimensions, format, mode
        """
        try:
            from PIL import Image

            with Image.open(path) as img:
                return {
                    "type": "image",
                    "format": img.format,
                    "mode": img.mode,
                    "width": img.width,
                    "height": img.height,
                    "size": f"{img.width}x{img.height}"
                }
        except Exception as e:
            logger.warning(f"[FileMetadataService] Image extraction error: {e}")
            return {"type": "image", "error": str(e)}

    def _safe_json_preview(self, data: Any, max_length: int = 500) -> Any:
        """
        Create a safe preview of JSON data (limited size).

        Args:
            data: JSON data to preview
            max_length: Maximum string length for preview

        Returns:
            Truncated or summarized data
        """
        json_str = json.dumps(data, indent=2)
        if len(json_str) > max_length:
            return f"{json_str[:max_length]}... (truncated)"
        return data

    def _calculate_json_depth(self, data: Any, current_depth: int = 0) -> int:
        """
        Calculate maximum nesting depth of JSON structure.

        Args:
            data: JSON data
            current_depth: Current recursion depth

        Returns:
            Maximum depth as integer
        """
        if isinstance(data, dict):
            if not data:
                return current_depth
            return max(self._calculate_json_depth(v, current_depth + 1) for v in data.values())
        elif isinstance(data, list):
            if not data:
                return current_depth
            return max(self._calculate_json_depth(item, current_depth + 1) for item in data)
        else:
            return current_depth

    def extract_multiple(self, file_paths: List[str]) -> Dict[str, Dict[str, Any]]:
        """
        Extract metadata from multiple files.

        Args:
            file_paths: List of file paths

        Returns:
            Dictionary mapping file paths to their metadata

        Example:
            >>> service = FileMetadataService()
            >>> results = service.extract_multiple(["/path/file1.csv", "/path/file2.json"])
            >>> for path, metadata in results.items():
            ...     print(f"{path}: {metadata['type']}")
        """
        results = {}
        for file_path in file_paths:
            results[file_path] = self.extract_metadata(file_path)
        return results

    def is_supported(self, file_path: str) -> bool:
        """
        Check if file format is supported.

        Args:
            file_path: Path to file

        Returns:
            True if format is supported, False otherwise
        """
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_formats


# Global service instance (optional, for convenience)
file_metadata_service = FileMetadataService()
