"""DOCX file handler with specialized metadata extraction."""

from pathlib import Path
from typing import Any, Dict

from .base import BaseFileHandler


class DOCXFileHandler(BaseFileHandler):
    """Handler for DOCX files."""

    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']

    def extract_metadata(
        self,
        file_path: Path,
        quick_mode: bool = False
    ) -> Dict[str, Any]:
        """Extract metadata from DOCX file."""
        metadata = {
            'file_type': 'docx',
            'file_size': self._get_file_size(file_path),
            'error': None
        }

        try:
            from docx import Document

            doc = Document(file_path)

            metadata['num_paragraphs'] = len(doc.paragraphs)
            metadata['num_tables'] = len(doc.tables)

            if not quick_mode:
                # Extract text preview
                text_preview = []
                for para in doc.paragraphs[:5]:  # First 5 paragraphs
                    if para.text.strip():
                        text_preview.append(para.text.strip())

                metadata['preview'] = '\n'.join(text_preview)[:500]

                # Table info
                if doc.tables:
                    table_info = []
                    for i, table in enumerate(doc.tables[:3]):  # First 3 tables
                        table_info.append({
                            'index': i,
                            'rows': len(table.rows),
                            'cols': len(table.columns) if table.rows else 0
                        })
                    metadata['table_info'] = table_info
            else:
                metadata['preview'] = 'Quick mode - preview not loaded'

        except Exception as e:
            metadata['error'] = str(e)

        return metadata

    def build_context_section(
        self,
        filename: str,
        metadata: Dict[str, Any],
        index: int
    ) -> str:
        """Build context section for DOCX file."""
        lines = []
        lines.append(f"{index}. {filename} (DOCX)")

        if metadata.get('error'):
            lines.append(f"   Error: {metadata['error']}")
            return '\n'.join(lines)

        # Basic info
        num_paragraphs = metadata.get('num_paragraphs', 0)
        num_tables = metadata.get('num_tables', 0)
        lines.append(f"   Paragraphs: {num_paragraphs}")
        lines.append(f"   Tables: {num_tables}")

        # Table details
        table_info = metadata.get('table_info', [])
        if table_info:
            lines.append("   Table structure:")
            for tbl in table_info:
                lines.append(f"     Table {tbl['index']}: {tbl['rows']} rows × {tbl['cols']} cols")

        # Preview
        preview = metadata.get('preview', '')
        if preview and preview != 'Quick mode - preview not loaded':
            preview_short = preview[:200]
            if len(preview) > 200:
                preview_short += '...'
            lines.append(f"   Preview: {preview_short}")

        return '\n'.join(lines)
