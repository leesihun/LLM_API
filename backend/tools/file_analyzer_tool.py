"""
File Analyzer Tool
==================
Analyzes uploaded files to extract metadata, format, structure, and preview.
This is a lightweight alternative to PYTHON_CODER for initial file inspection.

Version: 1.0.0
Created: 2025-01-05
"""

import os
import json
import logging
import uuid
from typing import Dict, Any, List, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class FileAnalyzer:
    """
    Analyzes files to extract metadata and structural information.
    Supports: CSV, Excel, JSON, TXT, PDF, Images, and more.

    Analysis Modes:
    - Standard: Fast direct analysis with recursive structure exploration (default)
    - Deep LLM: Uses python_coder_tool to generate custom analysis code (optional)
    """

    def __init__(self, use_llm_for_complex: bool = False):
        """
        Initialize FileAnalyzer.

        Args:
            use_llm_for_complex: If True, uses LLM-based analysis for complex structures
        """
        self.use_llm_for_complex = use_llm_for_complex
        self.supported_formats = {
            'csv': self._analyze_csv,
            'xlsx': self._analyze_excel,
            'xls': self._analyze_excel,
            'xlsm': self._analyze_excel,
            'json': self._analyze_json,
            'txt': self._analyze_text,
            'pdf': self._analyze_pdf,
            'docx': self._analyze_docx,
            'doc': self._analyze_doc,
            'png': self._analyze_image,
            'jpg': self._analyze_image,
            'jpeg': self._analyze_image,
            'gif': self._analyze_image,
            'bmp': self._analyze_image,
        }

    def analyze(self, file_paths: List[str], user_query: str = "") -> Dict[str, Any]:
        """
        Main analysis entry point.

        Args:
            file_paths: List of file paths to analyze
            user_query: User's question (optional, for context)

        Returns:
            Dict with analysis results
        """
        try:
            if not file_paths:
                return {
                    "success": False,
                    "error": "No files provided for analysis"
                }

            results = []
            for file_path in file_paths:
                if not os.path.exists(file_path):
                    results.append({
                        "file": file_path,
                        "success": False,
                        "error": f"File not found: {file_path}"
                    })
                    continue

                file_result = self._analyze_single_file(file_path)
                results.append(file_result)

            # Generate summary
            summary = self._generate_summary(results, user_query)

            return {
                "success": True,
                "files_analyzed": len(file_paths),
                "results": results,
                "summary": summary
            }

        except Exception as e:
            logger.error(f"File analysis error: {e}", exc_info=True)
            return {
                "success": False,
                "error": f"Analysis failed: {str(e)}"
            }

    def _analyze_single_file(self, file_path: str) -> Dict[str, Any]:
        """Analyze a single file."""
        try:
            path = Path(file_path)
            extension = path.suffix.lstrip('.').lower()

            # Basic file info
            file_info = {
                "file": str(path.name),
                "full_path": str(path.absolute()),
                "extension": extension,
                "size_bytes": os.path.getsize(file_path),
                "size_human": self._human_readable_size(os.path.getsize(file_path)),
                "success": True
            }

            # Format-specific analysis
            if extension in self.supported_formats:
                analyzer_func = self.supported_formats[extension]
                detailed_info = analyzer_func(file_path)
                file_info.update(detailed_info)
            else:
                file_info["format"] = "Unknown/Unsupported"
                file_info["note"] = f"Format '{extension}' not specifically supported"

            return file_info

        except Exception as e:
            return {
                "file": file_path,
                "success": False,
                "error": str(e)
            }

    def _analyze_csv(self, file_path: str) -> Dict[str, Any]:
        """Analyze CSV file with advanced delimiter detection and data quality profiling."""
        try:
            import pandas as pd
            import csv

            # Auto-detect delimiter
            delimiter = ','
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    sample = f.read(10240)  # Read first 10KB
                    sniffer = csv.Sniffer()
                    delimiter = sniffer.sniff(sample).delimiter
            except:
                pass  # Use default comma

            # Read with various encodings
            encodings = ['utf-8', 'cp949', 'euc-kr', 'latin1', 'iso-8859-1']
            df = None
            used_encoding = None

            for enc in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=enc, nrows=100, delimiter=delimiter)
                    used_encoding = enc
                    break
                except:
                    continue

            if df is None:
                return {"format": "CSV", "error": "Failed to read CSV with any encoding"}

            # Extract metadata
            full_df = pd.read_csv(file_path, encoding=used_encoding, delimiter=delimiter)

            # Data quality profiling
            null_counts = full_df.isnull().sum().to_dict()
            total_nulls = sum(null_counts.values())
            duplicate_rows = full_df.duplicated().sum()

            # Detect numeric columns with outliers
            numeric_cols = full_df.select_dtypes(include=['number']).columns.tolist()

            return {
                "format": "CSV",
                "encoding": used_encoding,
                "delimiter": repr(delimiter),  # Show special chars like \t
                "rows": len(full_df),
                "columns": len(full_df.columns),
                "column_names": list(full_df.columns),
                "column_types": {col: str(dtype) for col, dtype in full_df.dtypes.items()},
                "preview": full_df.head(5).to_dict(orient='records'),
                "null_counts": null_counts,
                "total_null_values": int(total_nulls),
                "null_percentage": round(total_nulls / (len(full_df) * len(full_df.columns)) * 100, 2) if len(full_df) > 0 else 0,
                "duplicate_rows": int(duplicate_rows),
                "numeric_columns": numeric_cols,
                "memory_usage": f"{full_df.memory_usage(deep=True).sum() / 1024:.2f} KB"
            }

        except ImportError:
            return {"format": "CSV", "error": "pandas not installed"}
        except Exception as e:
            return {"format": "CSV", "error": str(e)}

    def _analyze_excel(self, file_path: str) -> Dict[str, Any]:
        """Analyze Excel file with advanced metadata extraction."""
        try:
            import pandas as pd
            from openpyxl import load_workbook
            from openpyxl.utils import get_column_letter

            # Get all sheet names
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names

            # Load workbook for advanced metadata
            try:
                wb = load_workbook(file_path, data_only=False, read_only=True)
            except:
                wb = None

            sheets_info = []
            formulas_found = []
            named_ranges = []
            merged_cells_info = []

            # Analyze ALL sheets (not just first 3)
            for idx, sheet_name in enumerate(sheet_names):
                # Use pandas for data analysis
                df = pd.read_excel(file_path, sheet_name=sheet_name, nrows=100)
                full_df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Data quality metrics
                null_counts = full_df.isnull().sum().to_dict()
                total_nulls = sum(null_counts.values())
                duplicate_rows = full_df.duplicated().sum()

                sheet_info = {
                    "sheet_name": sheet_name,
                    "rows": len(full_df),
                    "columns": len(full_df.columns),
                    "column_names": list(full_df.columns),
                    "column_types": {col: str(dtype) for col, dtype in full_df.dtypes.items()},
                    "preview": df.head(3).to_dict(orient='records'),
                    "null_counts": null_counts,
                    "total_null_values": int(total_nulls),
                    "duplicate_rows": int(duplicate_rows)
                }

                # Extract advanced metadata using openpyxl
                if wb:
                    try:
                        ws = wb[sheet_name]

                        # Find formulas (sample first 100 cells)
                        formulas_in_sheet = []
                        for row in list(ws.iter_rows(max_row=min(100, ws.max_row), max_col=ws.max_column))[:20]:
                            for cell in row:
                                if cell.value and isinstance(cell.value, str) and cell.value.startswith('='):
                                    formulas_in_sheet.append({
                                        "cell": f"{get_column_letter(cell.column)}{cell.row}",
                                        "formula": str(cell.value)[:100]  # Truncate long formulas
                                    })

                        if formulas_in_sheet:
                            formulas_found.append({
                                "sheet": sheet_name,
                                "formulas": formulas_in_sheet[:10]  # Max 10 formulas per sheet
                            })

                        # Detect merged cells
                        if hasattr(ws, 'merged_cells') and ws.merged_cells:
                            merged_ranges = [str(merged_cell) for merged_cell in list(ws.merged_cells.ranges)[:10]]
                            if merged_ranges:
                                merged_cells_info.append({
                                    "sheet": sheet_name,
                                    "merged_ranges": merged_ranges
                                })

                        sheet_info["has_formulas"] = len(formulas_in_sheet) > 0
                        sheet_info["formula_count"] = len(formulas_in_sheet)
                        sheet_info["has_merged_cells"] = len(ws.merged_cells) > 0 if hasattr(ws, 'merged_cells') else False

                    except Exception as e:
                        logger.debug(f"Could not extract advanced metadata from sheet {sheet_name}: {e}")

                sheets_info.append(sheet_info)

            # Extract workbook-level metadata
            if wb:
                try:
                    # Named ranges
                    if hasattr(wb, 'defined_names'):
                        for named_range in wb.defined_names.definedName:
                            named_ranges.append({
                                "name": named_range.name,
                                "scope": named_range.localSheetId if named_range.localSheetId is not None else "Workbook",
                                "refers_to": str(named_range.attr_text)[:100]
                            })
                except Exception as e:
                    logger.debug(f"Could not extract named ranges: {e}")

            result = {
                "format": "Excel",
                "total_sheets": len(sheet_names),
                "sheet_names": sheet_names,
                "sheets_analyzed": sheets_info,
                "has_formulas": len(formulas_found) > 0,
                "formulas": formulas_found if formulas_found else None,
                "has_named_ranges": len(named_ranges) > 0,
                "named_ranges": named_ranges[:20] if named_ranges else None,  # Max 20
                "has_merged_cells": len(merged_cells_info) > 0,
                "merged_cells": merged_cells_info if merged_cells_info else None
            }

            return result

        except ImportError:
            return {"format": "Excel", "error": "pandas or openpyxl not installed"}
        except Exception as e:
            logger.error(f"Excel analysis error: {e}", exc_info=True)
            return {"format": "Excel", "error": str(e)}

    def _analyze_json(self, file_path: str) -> Dict[str, Any]:
        """Analyze JSON file with deep structure exploration."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            result = {
                "format": "JSON",
                "structure": type(data).__name__
            }

            # Basic structure info
            if isinstance(data, list):
                result["items_count"] = len(data)
                if len(data) > 0:
                    result["first_item_type"] = type(data[0]).__name__
                    result["preview"] = data[:]
                    if isinstance(data[0], dict):
                        result["keys"] = list(data[0].keys())
            elif isinstance(data, dict):
                result["keys"] = list(data.keys())
                result["preview"] = {k: data[k] for k in list(data.keys())[:]}
            else:
                result["value"] = str(data)[:]

            # Deep structure analysis (recursive exploration)
            structure_info = self._explore_json_structure(data)
            result["depth_analysis"] = structure_info

            # Add human-readable structure summary
            result["structure_summary"] = self._generate_structure_summary(structure_info)

            return result

        except Exception as e:
            return {"format": "JSON", "error": str(e)}

    def _explore_json_structure(self, obj: Any, current_depth: int = 0, max_depth: int = 10, path: str = "root") -> Dict[str, Any]:
        """
        Recursively explore JSON structure to find all nested levels.

        Args:
            obj: Current object to analyze
            current_depth: Current nesting depth
            max_depth: Maximum depth to explore
            path: Current path in the structure

        Returns:
            Dictionary with structure information
        """
        if current_depth > max_depth:
            return {
                "type": "max_depth_reached",
                "depth": current_depth,
                "path": path
            }

        if isinstance(obj, dict):
            keys = list(obj.keys())
            children = {}

            # Analyze ALL keys (no sampling - complete structure exploration)
            for key in keys:
                child_path = f"{path}.{key}"
                children[key] = self._explore_json_structure(obj[key], current_depth + 1, max_depth, child_path)

            return {
                "type": "dict",
                "depth": current_depth,
                "path": path,
                "keys": keys,
                "key_count": len(keys),
                "children": children
            }

        elif isinstance(obj, list):
            length = len(obj)
            sample_item = None

            # Analyze first item as representative
            if length > 0:
                sample_path = f"{path}[0]"
                sample_item = self._explore_json_structure(obj[0], current_depth + 1, max_depth, sample_path)

            return {
                "type": "list",
                "depth": current_depth,
                "path": path,
                "length": length,
                "sample_item": sample_item
            }

        else:
            # Leaf node (primitive type)
            return {
                "type": type(obj).__name__,
                "depth": current_depth,
                "path": path,
                "value_sample": str(obj)[:50] if obj is not None else None
            }

    def _generate_structure_summary(self, structure_info: Dict[str, Any]) -> str:
        """
        Generate human-readable summary of JSON structure.

        Args:
            structure_info: Structure information from _explore_json_structure

        Returns:
            Human-readable summary string
        """
        summary_lines = []
        max_depth = self._find_max_depth(structure_info)
        all_paths = self._collect_all_paths(structure_info)

        summary_lines.append(f"Maximum nesting depth: {max_depth}")
        summary_lines.append(f"Total unique paths: {len(all_paths)}")
        summary_lines.append("")
        summary_lines.append("Structure hierarchy:")

        # Show example paths at different depths
        paths_by_depth = {}
        for path_info in all_paths:
            depth = path_info["depth"]
            if depth not in paths_by_depth:
                paths_by_depth[depth] = []
            paths_by_depth[depth].append(path_info)

        for depth in sorted(paths_by_depth.keys()):
            paths = paths_by_depth[depth]
            summary_lines.append(f"  Depth {depth}: {len(paths)} node(s)")
            # Show ALL paths at each depth (no sampling)
            for path_info in paths:
                type_str = path_info["type"]
                path_str = path_info["path"]
                if type_str == "list":
                    summary_lines.append(f"    - {path_str} (array, {path_info.get('length', 0)} items)")
                elif type_str == "dict":
                    summary_lines.append(f"    - {path_str} (object, {path_info.get('key_count', 0)} keys)")
                else:
                    summary_lines.append(f"    - {path_str} ({type_str})")

        return "\n".join(summary_lines)

    def _find_max_depth(self, structure_info: Dict[str, Any]) -> int:
        """Find maximum depth in structure."""
        max_depth = structure_info.get("depth", 0)

        if structure_info.get("type") == "dict" and "children" in structure_info:
            for child in structure_info["children"].values():
                child_max = self._find_max_depth(child)
                max_depth = max(max_depth, child_max)
        elif structure_info.get("type") == "list" and "sample_item" in structure_info:
            if structure_info["sample_item"]:
                child_max = self._find_max_depth(structure_info["sample_item"])
                max_depth = max(max_depth, child_max)

        return max_depth

    def _collect_all_paths(self, structure_info: Dict[str, Any], paths: Optional[List[Dict[str, Any]]] = None) -> List[Dict[str, Any]]:
        """Collect all paths in the structure."""
        if paths is None:
            paths = []

        path_entry = {
            "path": structure_info.get("path", "root"),
            "type": structure_info.get("type", "unknown"),
            "depth": structure_info.get("depth", 0)
        }

        if structure_info.get("type") == "list":
            path_entry["length"] = structure_info.get("length", 0)
        elif structure_info.get("type") == "dict":
            path_entry["key_count"] = structure_info.get("key_count", 0)

        paths.append(path_entry)

        # Recurse into children
        if structure_info.get("type") == "dict" and "children" in structure_info:
            for child in structure_info["children"].values():
                self._collect_all_paths(child, paths)
        elif structure_info.get("type") == "list" and "sample_item" in structure_info:
            if structure_info["sample_item"]:
                self._collect_all_paths(structure_info["sample_item"], paths)

        return paths

    def _analyze_text(self, file_path: str) -> Dict[str, Any]:
        """Analyze text file."""
        try:
            encodings = ['utf-8', 'cp949', 'euc-kr', 'latin1']
            content = None
            used_encoding = None

            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        content = f.read()
                    used_encoding = enc
                    break
                except:
                    continue

            if content is None:
                return {"format": "Text", "error": "Failed to read text file"}

            lines = content.split('\n')

            return {
                "format": "Text",
                "encoding": used_encoding,
                "total_lines": len(lines),
                "total_characters": len(content),
                "total_words": len(content.split()),
                "preview": '\n'.join(lines[:10])
            }

        except Exception as e:
            return {"format": "Text", "error": str(e)}

    def _analyze_pdf(self, file_path: str) -> Dict[str, Any]:
        """Analyze PDF file."""
        try:
            import PyPDF2

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)

                # Extract text from first page
                first_page_text = pdf_reader.pages[0].extract_text() if num_pages > 0 else ""

                return {
                    "format": "PDF",
                    "total_pages": num_pages,
                    "first_page_preview": first_page_text[:500]
                }

        except ImportError:
            return {"format": "PDF", "error": "PyPDF2 not installed"}
        except Exception as e:
            return {"format": "PDF", "error": str(e)}

    def _analyze_image(self, file_path: str) -> Dict[str, Any]:
        """Analyze image file."""
        try:
            from PIL import Image

            with Image.open(file_path) as img:
                return {
                    "format": f"Image ({img.format})",
                    "dimensions": f"{img.width}x{img.height}",
                    "mode": img.mode,
                    "width": img.width,
                    "height": img.height
                }

        except ImportError:
            return {"format": "Image", "error": "PIL/Pillow not installed"}
        except Exception as e:
            return {"format": "Image", "error": str(e)}

    def _analyze_docx(self, file_path: str) -> Dict[str, Any]:
        """Analyze Word .docx file with comprehensive metadata extraction."""
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            total_text = '\n'.join(paragraphs)

            # Count words
            word_count = len(total_text.split())
            char_count = len(total_text)

            # Extract tables
            tables_info = []
            for idx, table in enumerate(doc.tables):
                rows = len(table.rows)
                cols = len(table.columns) if rows > 0 else 0

                # Extract table content (first 3 rows as sample)
                table_data = []
                for row_idx, row in enumerate(table.rows[:3]):
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                tables_info.append({
                    "table_number": idx + 1,
                    "rows": rows,
                    "columns": cols,
                    "preview": table_data
                })

            # Extract styles/headings
            headings = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    headings.append({
                        "level": para.style.name,
                        "text": para.text[:100]  # Truncate long headings
                    })

            # Extract images (count embedded media)
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1

            # Extract hyperlinks
            hyperlinks = []
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.underline or run.font.color:  # Likely a hyperlink
                        # Note: python-docx doesn't directly expose hyperlinks easily
                        # This is a simplified detection
                        pass

            # Detect formatting features
            has_bold = any(run.bold for para in doc.paragraphs for run in para.runs if run.bold)
            has_italic = any(run.italic for para in doc.paragraphs for run in para.runs if run.italic)
            has_underline = any(run.underline for para in doc.paragraphs for run in para.runs if run.underline)

            return {
                "format": "Word Document (.docx)",
                "total_paragraphs": len(paragraphs),
                "total_words": word_count,
                "total_characters": char_count,
                "total_tables": len(doc.tables),
                "total_images": image_count,
                "total_headings": len(headings),
                "headings": headings[:20] if headings else None,  # Max 20 headings
                "tables": tables_info if tables_info else None,
                "has_bold": has_bold,
                "has_italic": has_italic,
                "has_underline": has_underline,
                "text_preview": total_text[:500]  # First 500 chars
            }

        except ImportError:
            return {"format": "Word Document (.docx)", "error": "python-docx not installed"}
        except Exception as e:
            logger.error(f"DOCX analysis error: {e}", exc_info=True)
            return {"format": "Word Document (.docx)", "error": str(e)}

    def _analyze_doc(self, file_path: str) -> Dict[str, Any]:
        """Analyze legacy .doc file (limited support)."""
        try:
            # Try using textract if available
            try:
                import textract
                text = textract.process(file_path).decode('utf-8')

                return {
                    "format": "Word Document (.doc - legacy)",
                    "total_characters": len(text),
                    "total_words": len(text.split()),
                    "text_preview": text[:500],
                    "note": "Limited analysis for .doc format. Convert to .docx for full analysis."
                }
            except ImportError:
                # Fallback: try conversion via LibreOffice or suggest manual conversion
                return {
                    "format": "Word Document (.doc - legacy)",
                    "error": "textract not installed",
                    "note": "Please convert .doc to .docx format for full analysis. Install 'textract' for basic text extraction."
                }

        except Exception as e:
            logger.error(f"DOC analysis error: {e}", exc_info=True)
            return {"format": "Word Document (.doc - legacy)", "error": str(e)}

    def _human_readable_size(self, size_bytes: int) -> str:
        """Convert bytes to human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} TB"

    def _generate_summary(self, results: List[Dict[str, Any]], user_query: str) -> str:
        """Generate a human-readable summary of analysis results."""
        summary_lines = []

        successful = [r for r in results if r.get("success", False)]
        failed = [r for r in results if not r.get("success", False)]

        summary_lines.append(f"[File Analysis Summary] ({len(successful)}/{len(results)} successful)")
        summary_lines.append("")

        for result in successful:
            file_name = result.get("file", "Unknown")
            format_type = result.get("format", "Unknown")
            size = result.get("size_human", "Unknown size")

            summary_lines.append(f"- {file_name} ({format_type}, {size})")

            # Add format-specific details
            if format_type == "CSV":
                rows = result.get("rows", 0)
                cols = result.get("columns", 0)
                summary_lines.append(f"  - {rows:,} rows × {cols} columns")
                summary_lines.append(f"  - Columns: {', '.join(result.get('column_names', [])[:5])}")
                if result.get("duplicate_rows", 0) > 0:
                    summary_lines.append(f"  - Duplicate rows: {result.get('duplicate_rows')}")
                null_pct = result.get("null_percentage", 0)
                if null_pct > 0:
                    summary_lines.append(f"  - Missing data: {null_pct}%")

            elif format_type == "Excel":
                sheets = result.get("total_sheets", 0)
                summary_lines.append(f"  - {sheets} sheet(s): {', '.join(result.get('sheet_names', [])[:3])}")
                if result.get("has_formulas"):
                    summary_lines.append(f"  - Contains formulas")
                if result.get("has_named_ranges"):
                    summary_lines.append(f"  - Contains named ranges: {len(result.get('named_ranges', []))}")
                if result.get("has_merged_cells"):
                    summary_lines.append(f"  - Contains merged cells")

            elif format_type.startswith("Word Document"):
                words = result.get("total_words", 0)
                tables = result.get("total_tables", 0)
                images = result.get("total_images", 0)
                headings = result.get("total_headings", 0)
                summary_lines.append(f"  - {words:,} words, {tables} table(s), {images} image(s)")
                if headings > 0:
                    summary_lines.append(f"  - {headings} heading(s)")

            elif format_type == "JSON":
                structure = result.get("structure", "Unknown")
                summary_lines.append(f"  - Structure: {structure}")
                if "items_count" in result:
                    summary_lines.append(f"  - Items: {result['items_count']}")

            elif format_type == "Text":
                lines = result.get("total_lines", 0)
                words = result.get("total_words", 0)
                summary_lines.append(f"  - {lines:,} lines, {words:,} words")

            elif "Image" in format_type:
                dims = result.get("dimensions", "Unknown")
                summary_lines.append(f"  - Dimensions: {dims}")

            summary_lines.append("")

        if failed:
            summary_lines.append("[!] Failed files:")
            for result in failed:
                summary_lines.append(f"- {result.get('file', 'Unknown')}: {result.get('error', 'Unknown error')}")

        return "\n".join(summary_lines)

    def deep_analyze_with_llm(self, file_path: str, user_query: str = "") -> Dict[str, Any]:
        """
        Use LLM (via python_coder_tool) to generate and execute custom analysis code.
        This is more flexible but slower than direct analysis.

        Args:
            file_path: Path to file to analyze
            user_query: Specific user question to guide analysis

        Returns:
            Dict with analysis results from LLM-generated code
        """
        try:
            from backend.tools.python_coder_tool import python_coder_tool

            # Build analysis query for LLM
            analysis_prompt = f"""
Analyze the structure of this file in extreme detail:

1. Find the maximum nesting depth
2. Map out ALL key paths (e.g., data[0].user.profile.name)
3. Count items at each level
4. Show example values at leaf nodes
5. Identify all nested dictionaries and lists
6. Show the complete hierarchy

File: {os.path.basename(file_path)}

{f'User question: {user_query}' if user_query else ''}

Output a comprehensive JSON structure report.
"""

            # Execute analysis via python_coder_tool
            session_id = f"deep_analysis_{uuid.uuid4().hex[:8]}"

            result = python_coder_tool.execute_code_task(
                query=analysis_prompt,
                file_paths=[file_path],
                session_id=session_id
            )

            return {
                "format": "LLM Deep Analysis",
                "success": result.get("success", False),
                "analysis_result": result.get("result", ""),
                "code_generated": result.get("code", ""),
                "execution_time": result.get("execution_time", 0),
                "session_id": session_id
            }

        except ImportError:
            return {
                "format": "LLM Deep Analysis",
                "error": "python_coder_tool not available"
            }
        except Exception as e:
            logger.error(f"LLM deep analysis failed: {e}", exc_info=True)
            return {
                "format": "LLM Deep Analysis",
                "error": str(e)
            }


# Singleton instance
file_analyzer = FileAnalyzer()


def analyze_files(file_paths: List[str], user_query: str = "") -> Dict[str, Any]:
    """
    Convenience function to analyze files.

    Args:
        file_paths: List of file paths to analyze
        user_query: User's question (optional)

    Returns:
        Analysis results dictionary
    """
    return file_analyzer.analyze(file_paths, user_query)
