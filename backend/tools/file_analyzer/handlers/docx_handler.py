"""
DOCX File Handler
==================
Handler for analyzing Word .docx files with comprehensive metadata extraction.

Version: 1.0.0
Created: 2025-01-13
"""

from typing import Dict, Any, List

from backend.utils.logging_utils import get_logger
from ..base_handler import BaseFileHandler

logger = get_logger(__name__)


class DOCXHandler(BaseFileHandler):
    """
    Handler for DOCX file analysis.

    Features:
    - Paragraph and word counting
    - Table detection and extraction
    - Heading structure analysis
    - Image/media detection
    - Text formatting detection
    - Text preview extraction
    """

    def __init__(self):
        """Initialize DOCX handler."""
        self.supported_extensions = ['docx', 'doc']

    def supports(self, file_path: str) -> bool:
        """
        Check if this is a DOCX file.

        Args:
            file_path: Path to the file

        Returns:
            True if file has .docx or .doc extension
        """
        extension = self.get_file_extension(file_path)
        return extension in self.supported_extensions

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported extensions."""
        return self.supported_extensions

    def analyze(self, file_path: str) -> Dict[str, Any]:
        """
        Analyze Word document file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary with analysis results including:
            - format: 'Word Document (.docx)'
            - total_paragraphs: Number of paragraphs
            - total_words: Total word count
            - total_characters: Total character count
            - total_tables: Number of tables
            - total_images: Number of embedded images
            - total_headings: Number of headings
            - headings: List of heading structures
            - tables: Table structure information
            - has_bold/italic/underline: Formatting detection
            - text_preview: First 500 characters
        """
        extension = self.get_file_extension(file_path)

        # Handle legacy .doc files
        if extension == 'doc':
            return self._analyze_doc_legacy(file_path)

        # Handle .docx files
        try:
            from docx import Document

            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            total_text = '\n'.join(paragraphs)

            # Count words and characters
            word_count = len(total_text.split())
            char_count = len(total_text)

            # Extract tables
            tables_info = self._extract_tables(doc)

            # Extract styles/headings
            headings = self._extract_headings(doc)

            # Count embedded images
            image_count = self._count_images(doc)

            # Detect formatting features
            has_bold = any(run.bold for para in doc.paragraphs for run in para.runs if run.bold)
            has_italic = any(run.italic for para in doc.paragraphs for run in para.runs if run.italic)
            has_underline = any(run.underline for para in doc.paragraphs for run in para.runs if run.underline)

            return {
                "format": "Word Document (.docx)",
                "total_paragraphs": len(paragraphs),
                "total_words": word_count,
                "total_characters": char_count,
                "total_tables": len(doc.tables),
                "total_images": image_count,
                "total_headings": len(headings),
                "headings": headings[:20] if headings else None,
                "tables": tables_info if tables_info else None,
                "has_bold": has_bold,
                "has_italic": has_italic,
                "has_underline": has_underline,
                "text_preview": total_text[:500]
            }

        except ImportError:
            return {"format": "Word Document (.docx)", "error": "python-docx not installed"}
        except Exception as e:
            logger.error(f"DOCX analysis error: {e}", exc_info=True)
            return {"format": "Word Document (.docx)", "error": str(e)}

    def _extract_tables(self, doc) -> List[Dict[str, Any]]:
        """
        Extract table information from document.

        Args:
            doc: Document object

        Returns:
            List of table dictionaries
        """
        tables_info = []

        for idx, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns) if rows > 0 else 0

            # Extract table content (first 3 rows as sample)
            table_data = []
            for row in table.rows[:3]:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            tables_info.append({
                "table_number": idx + 1,
                "rows": rows,
                "columns": cols,
                "preview": table_data
            })

        return tables_info

    def _extract_headings(self, doc) -> List[Dict[str, str]]:
        """
        Extract heading structure from document.

        Args:
            doc: Document object

        Returns:
            List of heading dictionaries
        """
        headings = []

        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                headings.append({
                    "level": para.style.name,
                    "text": para.text  # Show full heading text
                })

        return headings

    def _count_images(self, doc) -> int:
        """
        Count embedded images in document.

        Args:
            doc: Document object

        Returns:
            Number of images
        """
        image_count = 0

        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
        except Exception:
            pass  # Image counting is optional

        return image_count

    def _analyze_doc_legacy(self, file_path: str) -> Dict[str, Any]:
        """
        Analyze legacy .doc file (limited support).

        Args:
            file_path: Path to .doc file

        Returns:
            Dictionary with limited analysis results
        """
        try:
            # Try using textract if available
            import textract
            text = textract.process(file_path).decode('utf-8')

            return {
                "format": "Word Document (.doc - legacy)",
                "total_characters": len(text),
                "total_words": len(text.split()),
                "text_preview": text[:500],
                "note": "Limited analysis for .doc format. Convert to .docx for full analysis."
            }

        except ImportError:
            return {
                "format": "Word Document (.doc - legacy)",
                "error": "textract not installed",
                "note": "Please convert .doc to .docx format for full analysis. Install 'textract' for basic text extraction."
            }
        except Exception as e:
            logger.error(f"DOC analysis error: {e}", exc_info=True)
            return {"format": "Word Document (.doc - legacy)", "error": str(e)}
